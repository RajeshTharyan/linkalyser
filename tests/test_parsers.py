from io import BytesIO

from docx import Document
from openpyxl import Workbook

from linkalyser.parsers import parse_excel, parse_html, parse_pdf, parse_word


def test_parse_html_strips_tags():
    text = parse_html(b"<html><body><h1>Title</h1><p>Hello world</p></body></html>")
    assert "Title" in text
    assert "Hello world" in text
    assert "<h1>" not in text


def test_parse_word_reads_paragraphs():
    document = Document()
    document.add_paragraph("climate policy")
    document.add_paragraph("second paragraph")
    buffer = BytesIO()
    document.save(buffer)
    text = parse_word(buffer.getvalue())
    assert "climate policy" in text
    assert "second paragraph" in text


def test_parse_excel_reads_cells():
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "funding"
    sheet["B1"] = 42
    buffer = BytesIO()
    workbook.save(buffer)
    text = parse_excel(buffer.getvalue())
    assert "funding" in text
    assert "42" in text


def test_corrupt_documents_return_empty_string():
    junk = b"this is not a document"
    assert parse_pdf(junk) == ""
    assert parse_word(junk) == ""
    assert parse_excel(junk) == ""
